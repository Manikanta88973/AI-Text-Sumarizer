import os
import docx
from PIL import Image, ImageDraw, ImageFont
from reportlab.pdfgen import canvas
from utils import (
    extract_from_txt,
    extract_from_docx,
    extract_from_pdf,
    extract_from_scanned_pdf,
    extract_from_handwritten_image,
    detect_and_extract,
    summarize_extracted_text
)

SAMPLES_DIR = os.path.join(os.path.dirname(__file__), "samples")
os.makedirs(SAMPLES_DIR, exist_ok=True)

def create_sample_files():
    """Create test samples for all supported file formats."""
    sample_text_content = (
        "Artificial intelligence and machine learning have revolutionized text processing. "
        "Modern neural network architectures allow automated summarization of long documents into concise key insights. "
        "This project extends input format support to include Word documents, PDFs, and handwritten notes. "
        "Optical character recognition and vision-encoder-decoder transformers make handwritten document extraction feasible."
    )

    # 1. Plain text file
    txt_path = os.path.join(SAMPLES_DIR, "sample.txt")
    with open(txt_path, "w", encoding="utf-8") as f:
        f.write(sample_text_content)

    # 2. Word .docx file
    docx_path = os.path.join(SAMPLES_DIR, "sample.docx")
    doc = docx.Document()
    doc.add_heading("Project Overview", level=1)
    doc.add_paragraph(sample_text_content)
    doc.save(docx_path)

    # 3. Vector text PDF file using ReportLab
    pdf_text_path = os.path.join(SAMPLES_DIR, "sample_text.pdf")
    c = canvas.Canvas(pdf_text_path)
    c.drawString(100, 750, "Sample Text-based PDF Document")
    c.drawString(100, 720, sample_text_content[:90])
    c.drawString(100, 700, sample_text_content[90:180])
    c.drawString(100, 680, sample_text_content[180:])
    c.save()

    # 4. Scanned PDF file (Image converted to PDF)
    scanned_pdf_path = os.path.join(SAMPLES_DIR, "sample_scanned.pdf")
    img_pdf = Image.new("RGB", (600, 300), color="white")
    d = ImageDraw.Draw(img_pdf)
    d.text((30, 50), "Scanned Invoice Document\nTotal Amount: $500\nPayment Status: Paid", fill="black")
    img_pdf.save(scanned_pdf_path, "PDF", resolution=100.0)

    # 5. Handwritten Note Image (.png)
    handwritten_img_path = os.path.join(SAMPLES_DIR, "sample_handwritten.png")
    img_hw = Image.new("RGB", (500, 100), color="white")
    d_hw = ImageDraw.Draw(img_hw)
    d_hw.text((20, 35), "hello world meeting at 5 pm tomorrow", fill="black")
    img_hw.save(handwritten_img_path)

    return {
        "txt": txt_path,
        "docx": docx_path,
        "pdf_text": pdf_text_path,
        "pdf_scanned": scanned_pdf_path,
        "handwritten_image": handwritten_img_path
    }

def run_tests():
    print("=" * 60)
    print("RUNNING MULTI-FORMAT EXTRACTION & SUMMARIZATION TESTS")
    print("=" * 60)

    samples = create_sample_files()

    for format_key, file_path in samples.items():
        print(f"\n--- Testing Format: [{format_key.upper()}] ({os.path.basename(file_path)}) ---")
        try:
            extracted_text, warning = detect_and_extract(file_path)
            print(f"Extracted Text Output:\n'{extracted_text}'")
            if warning:
                print(f"Warning Banner Output: {warning}")

            # Run summarization if text is sufficient length
            if len(extracted_text.strip()) >= 20:
                summary = summarize_extracted_text(extracted_text, max_length=100, min_length=15)
                print(f"Summarizer Output: '{summary}'")
            else:
                print("Text too short for summarization test.")

        except Exception as e:
            print(f"ERROR processing {format_key}: {e}")

    print("\n" + "=" * 60)
    print("ALL FORMAT EXTRACTION TESTS COMPLETED")
    print("=" * 60)

if __name__ == "__main__":
    run_tests()
