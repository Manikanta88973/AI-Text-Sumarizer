import os
import docx
import pdfplumber
import pytesseract
from PIL import Image
import torch

# Limit CPU threads to prevent memory spikes on constrained containers
torch.set_num_threads(1)

# Setup tesseract path if on Windows in default location
TESSERACT_PATHS = [
    r"C:\Program Files\Tesseract-OCR\tesseract.exe",
    r"C:\Program Files (x86)\Tesseract-OCR\tesseract.exe",
    os.path.expanduser(r"~\AppData\Local\Programs\Tesseract-OCR\tesseract.exe")
]
for p in TESSERACT_PATHS:
    if os.path.exists(p):
        pytesseract.pytesseract.tesseract_cmd = p
        break

# Global cached models (Lazy Loaded)
_trocr_processor = None
_trocr_model = None
_summarizer_tokenizer = None
_summarizer_model = None

def get_trocr_model():
    global _trocr_processor, _trocr_model
    if _trocr_processor is None or _trocr_model is None:
        from transformers import TrOCRProcessor, VisionEncoderDecoderModel
        model_name = "microsoft/trocr-base-handwritten"
        _trocr_processor = TrOCRProcessor.from_pretrained(model_name)
        _trocr_model = VisionEncoderDecoderModel.from_pretrained(model_name, low_cpu_mem_usage=True)
        _trocr_model.eval()
    return _trocr_processor, _trocr_model

def get_summarizer():
    global _summarizer_tokenizer, _summarizer_model
    if _summarizer_tokenizer is None or _summarizer_model is None:
        from transformers import AutoTokenizer, AutoModelForSeq2SeqLM
        model_name = "sshleifer/distilbart-cnn-12-6"
        _summarizer_tokenizer = AutoTokenizer.from_pretrained(model_name)
        _summarizer_model = AutoModelForSeq2SeqLM.from_pretrained(model_name, low_cpu_mem_usage=True)
        _summarizer_model.eval()
    return _summarizer_tokenizer, _summarizer_model

def extract_from_txt(file_path):
    """Extract text from plain text file (.txt)."""
    encodings = ['utf-8', 'latin-1', 'cp1252']
    for enc in encodings:
        try:
            with open(file_path, 'r', encoding=enc) as f:
                content = f.read().strip()
                if content:
                    return content
        except UnicodeDecodeError:
            continue
    with open(file_path, 'r', encoding='utf-8', errors='ignore') as f:
        return f.read().strip()

def extract_from_docx(file_path):
    """Extract all paragraph text from a Word document (.docx) using python-docx."""
    doc = docx.Document(file_path)
    paragraphs = [p.text.strip() for p in doc.paragraphs if p.text.strip()]
    return "\n\n".join(paragraphs)

def extract_from_pdf(file_path):
    """Extract text from a text-based PDF using pdfplumber."""
    text_content = []
    with pdfplumber.open(file_path) as pdf:
        for page in pdf.pages:
            page_text = page.extract_text()
            if page_text:
                text_content.append(page_text.strip())
    return "\n\n".join(text_content)

def extract_from_scanned_pdf(file_path):
    """Convert PDF pages to images via pypdfium2 (no Poppler binary required) and run OCR."""
    images = []
    # Try pypdfium2 first (pure Python wheel, no Poppler needed!)
    try:
        import pypdfium2 as pdfium
        pdf = pdfium.PdfDocument(file_path)
        for page in pdf:
            image = page.render(scale=2).to_pil()
            images.append(image)
    except Exception:
        # Fallback to pdf2image
        try:
            from pdf2image import convert_from_path
            images = convert_from_path(file_path)
        except Exception as e:
            raise RuntimeError(f"Could not convert PDF to images: {e}")

    extracted_pages = []
    for idx, img in enumerate(images):
        page_text = ""
        # Try pytesseract OCR
        try:
            page_text = pytesseract.image_to_string(img).strip()
        except Exception:
            page_text = ""

        # Fallback to TrOCR if pytesseract failed or returned empty
        if not page_text:
            try:
                processor, model = get_trocr_model()
                with torch.no_grad():
                    pv = processor(images=img.convert("RGB"), return_tensors="pt").pixel_values
                    g_ids = model.generate(pv)
                    page_text = processor.batch_decode(g_ids, skip_special_tokens=True)[0].strip()
            except Exception:
                pass

        if page_text:
            extracted_pages.append(page_text)

    return "\n\n".join(extracted_pages)

def extract_from_handwritten_image(image_path):
    """Use Hugging Face's microsoft/trocr-base-handwritten model to read handwritten text from an uploaded image."""
    image = Image.open(image_path).convert("RGB")
    width, height = image.size

    generated_text = ""
    # Try TrOCR model
    try:
        processor, model = get_trocr_model()
        with torch.no_grad():
            pixel_values = processor(images=image, return_tensors="pt").pixel_values
            generated_ids = model.generate(pixel_values)
            generated_text = processor.batch_decode(generated_ids, skip_special_tokens=True)[0]
    except Exception as e:
        print(f"TrOCR extraction notice: {e}")

    # If TrOCR single-pass text is very short and the image is tall, crop into horizontal line strips
    if len(generated_text.strip()) < 10 and height > 80:
        try:
            processor, model = get_trocr_model()
            line_height = min(60, height // 3)
            lines = []
            for top in range(0, height, line_height):
                bottom = min(top + line_height + 10, height)
                line_crop = image.crop((0, top, width, bottom))
                with torch.no_grad():
                    pv = processor(images=line_crop, return_tensors="pt").pixel_values
                    g_ids = model.generate(pv)
                    txt = processor.batch_decode(g_ids, skip_special_tokens=True)[0]
                    if txt.strip() and txt.strip() not in lines:
                        lines.append(txt.strip())
            if lines:
                generated_text = "\n".join(lines)
        except Exception:
            pass

    # Fallback to pytesseract if available and generated_text is still short
    if len(generated_text.strip()) < 10:
        try:
            tess_text = pytesseract.image_to_string(image).strip()
            if len(tess_text) > len(generated_text.strip()):
                generated_text = tess_text
        except Exception:
            pass

    return generated_text.strip()

def is_low_quality_text(text):
    """Check if extracted OCR text is low quality, near-empty, or garbage."""
    if not text or len(text.strip()) < 10:
        return True
    return False

def detect_and_extract(file_path):
    """Router function checking file extension/type and calling the right extractor."""
    if not file_path or not os.path.exists(file_path):
        raise FileNotFoundError(f"File not found: {file_path}")

    ext = os.path.splitext(file_path)[1].lower()
    warning_msg = None
    extracted_text = ""

    if ext == ".txt":
        extracted_text = extract_from_txt(file_path)

    elif ext == ".docx":
        extracted_text = extract_from_docx(file_path)

    elif ext == ".pdf":
        try:
            extracted_text = extract_from_pdf(file_path)
        except Exception:
            extracted_text = ""

        if not extracted_text or len(extracted_text.strip()) < 15:
            try:
                extracted_text = extract_from_scanned_pdf(file_path)
            except Exception as e:
                warning_msg = f"Scanned PDF extraction notice: {str(e)}"
                if not extracted_text:
                    extracted_text = f"Error extracting from PDF: {e}"

    elif ext in [".jpg", ".jpeg", ".png"]:
        try:
            extracted_text = extract_from_handwritten_image(file_path)
        except Exception as e:
            warning_msg = f"Image extraction notice: {str(e)}"
            try:
                img = Image.open(file_path)
                extracted_text = pytesseract.image_to_string(img).strip()
            except Exception as tess_err:
                extracted_text = f"Could not extract readable text from image. Details: {e}"

    else:
        raise ValueError(f"Unsupported file format: {ext}. Supported formats: .txt, .docx, .pdf, .jpg, .jpeg, .png")

    if is_low_quality_text(extracted_text) and not warning_msg:
        warning_msg = "Warning: The extracted text is very short or empty. Please ensure the document or image contains readable text."

    return extracted_text, warning_msg

def fallback_extractive_summarize(text, num_sentences=3):
    """Fallback extractive summarizer based on sentence scoring."""
    import re
    sentences = [s.strip() for s in re.split(r'(?<=[.!?]) +', text) if s.strip()]
    if len(sentences) <= num_sentences:
        return text

    words = [w.lower() for w in re.findall(r'\w+', text)]
    word_freq = {}
    for w in words:
        word_freq[w] = word_freq.get(w, 0) + 1

    sentence_scores = {}
    for s in sentences:
        s_words = [w.lower() for w in re.findall(r'\w+', s)]
        score = sum(word_freq.get(w, 0) for w in s_words)
        sentence_scores[s] = score / max(1, len(s_words))

    top_sentences = sorted(sentences, key=lambda s: sentence_scores[s], reverse=True)[:num_sentences]
    top_sentences_in_order = [s for s in sentences if s in top_sentences]
    return " ".join(top_sentences_in_order)

def summarize_extracted_text(text, max_length=150, min_length=30):
    """Generate summary from extracted text with guaranteed fallback."""
    if not text or len(text.strip()) < 15:
        return "Cannot summarize: Provided text is too short or empty."

    try:
        tokenizer, model = get_summarizer()
        inputs = tokenizer(text, max_length=1024, truncation=True, return_tensors="pt")
        
        input_len = len(inputs["input_ids"][0])
        effective_max = min(max_length, max(min_length + 10, input_len))
        effective_min = min(min_length, max(5, effective_max - 10))

        with torch.no_grad():
            summary_ids = model.generate(
                inputs["input_ids"],
                max_length=effective_max,
                min_length=effective_min,
                num_beams=2,
                early_stopping=True
            )
        summary = tokenizer.decode(summary_ids[0], skip_special_tokens=True)
        if summary.strip():
            return summary.strip()
    except Exception as e:
        print(f"Neural summarizer warning: {e}. Using extractive fallback.")

    return fallback_extractive_summarize(text)
